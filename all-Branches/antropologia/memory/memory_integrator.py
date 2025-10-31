#!/usr/bin/env python3
# ==============================================================================
# INTEGRADOR COMPLETO DEL SISTEMA DE MEMORIA SHEILY - VERSIÓN EXTENDIDA
# ==============================================================================

import hashlib
import json
import mimetypes
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# Procesadores básicos
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import chardet

    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

# Procesadores Excel
try:
    import pandas as pd

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# Procesadores de imágenes con OCR
try:
    from PIL import Image

    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pytesseract

    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

# Procesadores de audio
try:
    import speech_recognition as sr

    SPEECH_AVAILABLE = True
except ImportError:
    SPEECH_AVAILABLE = False

# Procesadores de video
try:
    import cv2

    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False

# Procesadores de archivos comprimidos
try:
    import tarfile
    import zipfile

    ARCHIVE_AVAILABLE = True
except ImportError:
    ARCHIVE_AVAILABLE = False


class AdvancedFileProcessor:
    """Procesador avanzado de múltiples tipos de archivos"""

    @staticmethod
    def detect_encoding(file_path: Path) -> str:
        """Detectar encoding del archivo"""
        if CHARDET_AVAILABLE:
            try:
                with open(file_path, "rb") as f:
                    result = chardet.detect(f.read(100000))
                    return result["encoding"] or "utf-8"
            except:
                pass
        return "utf-8"

    @staticmethod
    def read_text_file(file_path: Path) -> str:
        """Leer archivo de texto plano"""
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

        if CHARDET_AVAILABLE:
            detected = AdvancedFileProcessor.detect_encoding(file_path)
            encodings.insert(0, detected)

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def read_pdf(file_path: Path) -> str:
        """Leer archivo PDF con extracción de texto"""
        if not PDF_AVAILABLE:
            return f"[PDF - Instalar: pip install PyPDF2]\nArchivo: {file_path.name}"

        try:
            text = []
            metadata = []

            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)

                # Metadatos del PDF
                info = pdf_reader.metadata
                if info:
                    metadata.append("=== METADATOS PDF ===")
                    for key, value in info.items():
                        metadata.append(f"{key}: {value}")
                    metadata.append("")

                # Contenido
                metadata.append(f"Total de páginas: {len(pdf_reader.pages)}\n")

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text.append(f"--- Página {page_num} ---\n{page_text}")

            result = "\n".join(metadata) + "\n\n" + "\n\n".join(text)
            return result if result.strip() else "[PDF sin texto extraíble]"
        except Exception as e:
            return f"[Error leyendo PDF: {str(e)}]"

    @staticmethod
    def read_docx(file_path: Path) -> str:
        """Leer archivo DOCX con estructura"""
        if not DOCX_AVAILABLE:
            return f"[DOCX - Instalar: pip install python-docx]\nArchivo: {file_path.name}"

        try:
            doc = DocxDocument(file_path)
            content = []

            # Propiedades del documento
            core_props = doc.core_properties
            content.append("=== PROPIEDADES DOCUMENTO ===")
            content.append(f"Título: {core_props.title or 'Sin título'}")
            content.append(f"Autor: {core_props.author or 'Desconocido'}")
            content.append(f"Fecha: {core_props.created or 'N/A'}")
            content.append("")

            # Contenido
            content.append("=== CONTENIDO ===")
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            # Tablas
            if doc.tables:
                content.append("\n=== TABLAS ===")
                for i, table in enumerate(doc.tables, 1):
                    content.append(f"\nTabla {i}:")
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        content.append(row_text)

            return "\n".join(content) if content else "[DOCX vacío]"
        except Exception as e:
            return f"[Error leyendo DOCX: {str(e)}]"

    @staticmethod
    def read_excel(file_path: Path) -> str:
        """Leer archivos Excel (XLS, XLSX)"""
        if not PANDAS_AVAILABLE:
            return f"[Excel - Instalar: pip install pandas openpyxl]\nArchivo: {file_path.name}"

        try:
            content = []

            # Leer todas las hojas
            excel_file = pd.ExcelFile(file_path)
            content.append(f"=== ARCHIVO EXCEL: {file_path.name} ===")
            content.append(f"Hojas encontradas: {len(excel_file.sheet_names)}\n")

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                content.append(f"--- HOJA: {sheet_name} ---")
                content.append(f"Dimensiones: {df.shape[0]} filas x {df.shape[1]} columnas")
                content.append(f"Columnas: {', '.join(df.columns.astype(str))}")
                content.append("")

                # Vista previa de datos (primeras 10 filas)
                content.append("Vista previa:")
                preview = df.head(10).to_string()
                content.append(preview)

                # Estadísticas básicas para columnas numéricas
                numeric_cols = df.select_dtypes(include=["number"]).columns
                if len(numeric_cols) > 0:
                    content.append("\nEstadísticas:")
                    stats = df[numeric_cols].describe().to_string()
                    content.append(stats)

                content.append("\n" + "=" * 50 + "\n")

            return "\n".join(content)
        except Exception as e:
            return f"[Error leyendo Excel: {str(e)}]"

    @staticmethod
    def read_csv(file_path: Path) -> str:
        """Leer archivos CSV con análisis"""
        if not PANDAS_AVAILABLE:
            # Fallback a lectura básica
            return AdvancedFileProcessor.read_text_file(file_path)

        try:
            # Intentar detectar el delimitador
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                sample = f.read(4096)

            # Detectar delimitador común
            delimiters = [",", ";", "\t", "|"]
            delimiter = ","
            max_count = 0
            for d in delimiters:
                count = sample.count(d)
                if count > max_count:
                    max_count = count
                    delimiter = d

            df = pd.read_csv(file_path, delimiter=delimiter, encoding="utf-8", on_bad_lines="skip")

            content = []
            content.append(f"=== ARCHIVO CSV: {file_path.name} ===")
            content.append(f"Dimensiones: {df.shape[0]} filas x {df.shape[1]} columnas")
            content.append(f"Delimitador detectado: '{delimiter}'")
            content.append(f"Columnas: {', '.join(df.columns.astype(str))}")
            content.append("")

            # Vista previa
            content.append("Vista previa (primeras 15 filas):")
            content.append(df.head(15).to_string())

            # Información de tipos de datos
            content.append("\nTipos de datos:")
            content.append(df.dtypes.to_string())

            # Valores faltantes
            missing = df.isnull().sum()
            if missing.sum() > 0:
                content.append("\nValores faltantes:")
                content.append(missing[missing > 0].to_string())

            return "\n".join(content)
        except Exception as e:
            return f"[Error leyendo CSV: {str(e)}]"

    @staticmethod
    def read_image_with_ocr(file_path: Path) -> str:
        """Leer texto de imágenes con OCR"""
        if not PIL_AVAILABLE or not TESSERACT_AVAILABLE:
            return f"[OCR - Instalar: pip install pillow pytesseract]\nArchivo: {file_path.name}"

        try:
            image = Image.open(file_path)

            content = []
            content.append(f"=== IMAGEN: {file_path.name} ===")
            content.append(f"Formato: {image.format}")
            content.append(f"Dimensiones: {image.size[0]}x{image.size[1]} píxeles")
            content.append(f"Modo: {image.mode}")
            content.append("")

            # Realizar OCR
            content.append("=== TEXTO EXTRAÍDO (OCR) ===")
            text = pytesseract.image_to_string(image, lang="spa+eng")

            if text.strip():
                content.append(text)
            else:
                content.append("[No se detectó texto en la imagen]")

            return "\n".join(content)
        except Exception as e:
            return f"[Error procesando imagen: {str(e)}]"

    @staticmethod
    def read_audio(file_path: Path) -> str:
        """Transcribir audio a texto"""
        if not SPEECH_AVAILABLE:
            return f"[Audio - Instalar: pip install SpeechRecognition pydub]\nArchivo: {file_path.name}"

        try:
            recognizer = sr.Recognizer()

            content = []
            content.append(f"=== AUDIO: {file_path.name} ===")

            # Intentar transcribir
            with sr.AudioFile(str(file_path)) as source:
                audio_data = recognizer.record(source)

                content.append("Transcribiendo audio...")
                content.append("")

                try:
                    # Intentar con Google Speech Recognition
                    text = recognizer.recognize_google(audio_data, language="es-ES")
                    content.append("=== TRANSCRIPCIÓN ===")
                    content.append(text)
                except sr.UnknownValueError:
                    content.append("[No se pudo entender el audio]")
                except sr.RequestError as e:
                    content.append(f"[Error del servicio de transcripción: {e}]")

            return "\n".join(content)
        except Exception as e:
            return f"[Error procesando audio: {str(e)}]"

    @staticmethod
    def read_video(file_path: Path) -> str:
        """Extraer información de video"""
        if not CV2_AVAILABLE:
            return f"[Video - Instalar: pip install opencv-python]\nArchivo: {file_path.name}"

        try:
            video = cv2.VideoCapture(str(file_path))

            content = []
            content.append(f"=== VIDEO: {file_path.name} ===")

            # Propiedades del video
            fps = video.get(cv2.CAP_PROP_FPS)
            frame_count = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(video.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(video.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0

            content.append(f"Resolución: {width}x{height}")
            content.append(f"FPS: {fps:.2f}")
            content.append(f"Duración: {duration:.2f} segundos ({duration/60:.2f} minutos)")
            content.append(f"Total de frames: {frame_count}")

            video.release()

            content.append("\n[Nota: Para transcripción de video, usar herramientas especializadas]")

            return "\n".join(content)
        except Exception as e:
            return f"[Error procesando video: {str(e)}]"

    @staticmethod
    def read_json(file_path: Path) -> str:
        """Leer y formatear archivo JSON"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            content = []
            content.append(f"=== JSON: {file_path.name} ===")
            content.append(f"Tipo raíz: {type(data).__name__}")

            if isinstance(data, dict):
                content.append(f"Claves principales: {', '.join(list(data.keys())[:10])}")
            elif isinstance(data, list):
                content.append(f"Elementos: {len(data)}")

            content.append("\n=== CONTENIDO ===")
            formatted = json.dumps(data, indent=2, ensure_ascii=False)

            # Limitar tamaño si es muy grande
            if len(formatted) > 50000:
                content.append(formatted[:50000])
                content.append("\n... [JSON truncado por tamaño]")
            else:
                content.append(formatted)

            return "\n".join(content)
        except Exception as e:
            return f"[Error leyendo JSON: {str(e)}]"

    @staticmethod
    def read_archive(file_path: Path) -> str:
        """Leer archivos comprimidos y listar contenido"""
        if not ARCHIVE_AVAILABLE:
            return f"[Archivos - Soporte nativo de Python]\nArchivo: {file_path.name}"

        try:
            content = []
            content.append(f"=== ARCHIVO COMPRIMIDO: {file_path.name} ===")

            ext = file_path.suffix.lower()

            if ext == ".zip":
                with zipfile.ZipFile(file_path, "r") as zf:
                    content.append(f"Tipo: ZIP")
                    content.append(f"Archivos: {len(zf.namelist())}")
                    content.append("\n=== CONTENIDO ===")

                    for info in zf.infolist()[:100]:  # Limitar a 100 archivos
                        size_mb = info.file_size / (1024 * 1024)
                        content.append(f"  {info.filename} - {size_mb:.2f} MB")

            elif ext in [".tar", ".tar.gz", ".tgz", ".tar.bz2"]:
                with tarfile.open(file_path, "r") as tf:
                    members = tf.getmembers()
                    content.append(f"Tipo: TAR")
                    content.append(f"Archivos: {len(members)}")
                    content.append("\n=== CONTENIDO ===")

                    for member in members[:100]:
                        size_mb = member.size / (1024 * 1024)
                        content.append(f"  {member.name} - {size_mb:.2f} MB")

            return "\n".join(content)
        except Exception as e:
            return f"[Error leyendo archivo: {str(e)}]"

    @staticmethod
    def process_file(file_path: Path) -> Tuple[str, bool]:
        """Procesar archivo según su tipo - VERSIÓN EXTENDIDA"""
        ext = file_path.suffix.lower()

        try:
            # Documentos
            if ext == ".pdf":
                content = AdvancedFileProcessor.read_pdf(file_path)
                return content, True
            elif ext == ".docx":
                content = AdvancedFileProcessor.read_docx(file_path)
                return content, True

            # Hojas de cálculo
            elif ext in [".xlsx", ".xls"]:
                content = AdvancedFileProcessor.read_excel(file_path)
                return content, True
            elif ext == ".csv":
                content = AdvancedFileProcessor.read_csv(file_path)
                return content, True

            # Imágenes con OCR
            elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"]:
                content = AdvancedFileProcessor.read_image_with_ocr(file_path)
                return content, True

            # Audio
            elif ext in [".wav", ".mp3", ".ogg", ".flac"]:
                content = AdvancedFileProcessor.read_audio(file_path)
                return content, True

            # Video
            elif ext in [".mp4", ".avi", ".mkv", ".mov", ".wmv"]:
                content = AdvancedFileProcessor.read_video(file_path)
                return content, True

            # Archivos comprimidos
            elif ext in [".zip", ".tar", ".tar.gz", ".tgz", ".tar.bz2"]:
                content = AdvancedFileProcessor.read_archive(file_path)
                return content, True

            # JSON
            elif ext == ".json":
                content = AdvancedFileProcessor.read_json(file_path)
                return content, True

            # Archivos de texto y código
            elif ext in [
                ".txt",
                ".md",
                ".py",
                ".js",
                ".java",
                ".cpp",
                ".c",
                ".css",
                ".html",
                ".xml",
                ".sql",
                ".sh",
                ".bat",
                ".ps1",
                ".yml",
                ".yaml",
                ".toml",
                ".ini",
                ".cfg",
                ".conf",
            ]:
                content = AdvancedFileProcessor.read_text_file(file_path)
                return content, True

            else:
                return f"[Tipo de archivo {ext} no soportado]\nArchivo: {file_path.name}", False

        except Exception as e:
            return f"[Error procesando archivo: {str(e)}]", False


class SheilyMemoryIntegrator:
    """Integrador completo del sistema de memoria - VERSIÓN EXTENDIDA"""

    def __init__(self):
        self.memory_root = Path("/home/yo/Escritorio/Sheily-Final/sheily_core/memory")
        self.input_dir = Path("/home/yo/Escritorio/Sheily-Final/memoria_entrada/pendientes")
        self.processed_dir = Path("/home/yo/Escritorio/Sheily-Final/memoria_entrada/completados")

        # Crear directorios
        self.memory_root.mkdir(parents=True, exist_ok=True)
        self.input_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)

        # Subdirectorios de memoria
        self.memory_dirs = {
            "raw": self.memory_root / "raw_memories",
            "processed": self.memory_root / "processed_memories",
            "index": self.memory_root / "index",
            "metadata": self.memory_root / "metadata",
            "multimedia": self.memory_root / "multimedia_metadata",
        }

        for dir_path in self.memory_dirs.values():
            dir_path.mkdir(parents=True, exist_ok=True)

        self.config = {
            "max_file_size_mb": 100,  # Aumentado para multimedia
            "supported_extensions": [
                # Documentos
                ".txt",
                ".md",
                ".pdf",
                ".docx",
                ".html",
                ".json",
                ".xml",
                # Hojas de cálculo
                ".csv",
                ".xlsx",
                ".xls",
                # Código
                ".py",
                ".js",
                ".java",
                ".cpp",
                ".c",
                ".css",
                ".sql",
                ".sh",
                ".bat",
                ".ps1",
                ".yml",
                ".yaml",
                ".toml",
                ".ini",
                ".cfg",
                ".conf",
                # Imágenes
                ".png",
                ".jpg",
                ".jpeg",
                ".bmp",
                ".tiff",
                ".gif",
                # Audio
                ".wav",
                ".mp3",
                ".ogg",
                ".flac",
                # Video
                ".mp4",
                ".avi",
                ".mkv",
                ".mov",
                ".wmv",
                # Archivos
                ".zip",
                ".tar",
                ".tar.gz",
                ".tgz",
                ".tar.bz2",
            ],
            "chunk_size": 1000,
            "overlap_size": 200,
        }

        self.memory_index_file = self.memory_dirs["index"] / "memory_index.json"
        self.memory_index = self.load_memory_index()

    def load_memory_index(self) -> Dict[str, Any]:
        """Cargar índice de memorias existente"""
        if self.memory_index_file.exists():
            try:
                with open(self.memory_index_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                pass
        return {
            "created": datetime.now().isoformat(),
            "total_memories": 0,
            "memories": {},
            "categories": {},
            "tags": {},
            "file_types": {},
        }

    def save_memory_index(self):
        """Guardar índice de memorias"""
        self.memory_index["last_updated"] = datetime.now().isoformat()
        with open(self.memory_index_file, "w", encoding="utf-8") as f:
            json.dump(self.memory_index, f, indent=2, ensure_ascii=False)

    def generate_memory_id(self, file_path: Path) -> str:
        """Generar ID único para la memoria"""
        timestamp = datetime.now().isoformat()
        unique_string = f"{file_path.name}_{timestamp}_{file_path.stat().st_size}"
        return hashlib.sha256(unique_string.encode()).hexdigest()[:16]

    def scan_input_directory(self) -> List[Dict[str, Any]]:
        """Escanear directorio de entrada"""
        files_info = []

        if not self.input_dir.exists():
            return files_info

        for file_path in self.input_dir.iterdir():
            if file_path.is_file() and not file_path.name.startswith("."):
                file_info = {
                    "file_name": file_path.name,
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size,
                    "file_size_mb": file_path.stat().st_size / (1024 * 1024),
                    "extension": file_path.suffix.lower(),
                    "modified_time": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                    "status": "pending",
                }
                files_info.append(file_info)

        return files_info

    def validate_file(self, file_info: Dict[str, Any]) -> Tuple[bool, str]:
        """Validar archivo"""
        if file_info["extension"] not in self.config["supported_extensions"]:
            return False, f"Extensión {file_info['extension']} no soportada"

        if file_info["file_size_mb"] > self.config["max_file_size_mb"]:
            return (
                False,
                f"Archivo demasiado grande ({file_info['file_size_mb']:.1f}MB > {self.config['max_file_size_mb']}MB)",
            )

        return True, "Archivo válido"

    def create_chunks(self, content: str) -> List[Dict[str, Any]]:
        """Crear chunks del contenido"""
        chunks = []
        chunk_size = self.config["chunk_size"]
        overlap = self.config["overlap_size"]

        if len(content) <= chunk_size:
            chunks.append(
                {
                    "chunk_id": "chunk_001",
                    "content": content,
                    "start_pos": 0,
                    "end_pos": len(content),
                    "size": len(content),
                }
            )
        else:
            start_pos = 0
            chunk_id = 1

            while start_pos < len(content):
                end_pos = min(start_pos + chunk_size, len(content))
                chunk_content = content[start:end_pos]

                chunks.append(
                    {
                        "chunk_id": f"chunk_{chunk_id:03d}",
                        "content": chunk_content,
                        "start_pos": start_pos,
                        "end_pos": end_pos,
                        "size": len(chunk_content),
                    }
                )

                start_pos += chunk_size - overlap
                chunk_id += 1

                if end_pos >= len(content):
                    break

        return chunks

    def generate_summary(self, content: str) -> str:
        """Generar resumen automático"""
        lines = [line.strip() for line in content.split("\n") if line.strip()]

        if not lines:
            return "Contenido vacío"

        summary_lines = []
        for line in lines[:10]:
            if len(line) > 20:
                summary_lines.append(line)
            if len(summary_lines) >= 3:
                break

        summary = " ".join(summary_lines)

        if len(summary) > 300:
            summary = summary[:297] + "..."

        return summary if summary else "Sin resumen disponible"

    def calculate_importance(self, content: str) -> float:
        """Calcular score de importancia"""
        score = 0.0

        if len(content) > 5000:
            score += 0.3
        if content.count("\n") > 10:
            score += 0.2

        importance_keywords = [
            "importante",
            "crucial",
            "esencial",
            "fundamental",
            "crítico",
            "important",
            "critical",
            "essential",
            "key",
            "vital",
        ]

        if any(keyword in content.lower() for keyword in importance_keywords):
            score += 0.3

        if content.count("#") > 3 or content.count("-") > 10:
            score += 0.2

        return min(score, 1.0)

    def detect_category(self, content: str, file_name: str) -> str:
        """Detectar categoría del contenido"""
        content_lower = content.lower()
        file_lower = file_name.lower()

        categories = {
            "código": ["def ", "class ", "function", "import ", "const ", "var "],
            "académico": ["universidad", "estudio", "investigación", "tesis"],
            "técnico": ["configuración", "instalación", "comando", "script"],
            "documentación": ["documento", "guía", "manual", "tutorial"],
            "datos": ["csv", "tabla", "datos", "excel", "hoja de cálculo"],
            "multimedia": ["imagen", "audio", "video", "transcripción"],
            "literario": ["novela", "cuento", "poesía", "capítulo"],
            "científico": ["experimento", "hipótesis", "método", "resultado"],
        }

        for category, keywords in categories.items():
            matches = sum(1 for keyword in keywords if keyword in content_lower or keyword in file_lower)
            if matches >= 2:
                return category

        return "general"

    def extract_tags(self, content: str, file_name: str) -> List[str]:
        """Extraer tags relevantes"""
        tags = []

        file_words = re.findall(r"\w+", file_name.lower())
        file_tags = [word for word in file_words if len(word) > 3 and word.isalpha()]
        tags.extend(file_tags[:3])

        words = re.findall(r"\b[a-záéíóúñ]{5,}\b", content.lower())
        word_freq = {}
        for word in words:
            if word.isalpha():
                word_freq[word] = word_freq.get(word, 0) + 1

        frequent_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:8]
        content_tags = [word for word, freq in frequent_words if freq > 2]
        tags.extend(content_tags[:4])

        return list(set(tags))[:6]

    def create_memory_entry(self, file_info: Dict[str, Any], content: str, memory_id: str) -> Dict[str, Any]:
        """Crear entrada de memoria estructurada"""
        chunks = self.create_chunks(content)
        category = self.detect_category(content, file_info["file_name"])
        tags = self.extract_tags(content, file_info["file_name"])

        memory_entry = {
            "memory_id": memory_id,
            "metadata": {
                "source_file": file_info["file_name"],
                "file_size": file_info["file_size"],
                "file_type": file_info["extension"],
                "processing_time": datetime.now().isoformat(),
                "total_chunks": len(chunks),
                "total_characters": len(content),
                "memory_type": "documento_procesado",
            },
            "content": {
                "full_text": content,
                "chunks": chunks,
                "summary": self.generate_summary(content),
            },
            "context": {
                "importance_score": self.calculate_importance(content),
                "category": category,
                "tags": tags,
            },
            "access_info": {
                "created": datetime.now().isoformat(),
                "last_accessed": None,
                "access_count": 0,
                "retrieval_score": 0.0,
            },
        }

        return memory_entry

    def save_memory(self, memory_entry: Dict[str, Any]) -> bool:
        """Guardar memoria en el sistema"""
        try:
            memory_id = memory_entry["memory_id"]

            # Guardar memoria completa
            memory_file = self.memory_dirs["processed"] / f"{memory_id}.json"
            with open(memory_file, "w", encoding="utf-8") as f:
                json.dump(memory_entry, f, indent=2, ensure_ascii=False)

            # Actualizar índice
            self.memory_index["memories"][memory_id] = {
                "file_name": memory_entry["metadata"]["source_file"],
                "category": memory_entry["context"]["category"],
                "tags": memory_entry["context"]["tags"],
                "created": memory_entry["access_info"]["created"],
                "importance": memory_entry["context"]["importance_score"],
                "file_type": memory_entry["metadata"]["file_type"],
            }

            # Actualizar categorías
            category = memory_entry["context"]["category"]
            if category not in self.memory_index["categories"]:
                self.memory_index["categories"][category] = []
            self.memory_index["categories"][category].append(memory_id)

            # Actualizar tags
            for tag in memory_entry["context"]["tags"]:
                if tag not in self.memory_index["tags"]:
                    self.memory_index["tags"][tag] = []
                self.memory_index["tags"][tag].append(memory_id)

            # Actualizar tipos de archivo
            file_type = memory_entry["metadata"]["file_type"]
            if file_type not in self.memory_index["file_types"]:
                self.memory_index["file_types"][file_type] = []
            self.memory_index["file_types"][file_type].append(memory_id)

            self.memory_index["total_memories"] = len(self.memory_index["memories"])
            self.save_memory_index()

            return True

        except Exception as e:
            print(f"❌ Error guardando memoria: {e}")
            return False

    def process_memory_integration(self) -> Dict[str, Any]:
        """Procesar integración completa de memoria"""
        print("🧠 SISTEMA DE MEMORIA SHEILY - PROCESAMIENTO EXTENDIDO")
        print("=" * 70)

        input_files = self.scan_input_directory()

        if not input_files:
            return {
                "status": "no_files",
                "message": "No hay archivos en memoria_entrada/pendientes/",
                "files_found": 0,
            }

        print(f"📂 Archivos encontrados: {len(input_files)}\n")

        results = {
            "timestamp": datetime.now().isoformat(),
            "total_files": len(input_files),
            "processed_files": 0,
            "successful_integrations": 0,
            "failed_integrations": 0,
            "files": [],
        }

        for file_info in input_files:
            print(f"📄 Procesando: {file_info['file_name']}")
            print(f"   Tamaño: {file_info['file_size_mb']:.2f} MB | Tipo: {file_info['extension']}")

            # Validar archivo
            is_valid, validation_message = self.validate_file(file_info)

            if not is_valid:
                print(f"   ❌ {validation_message}\n")
                file_info["status"] = "rejected"
                file_info["rejection_reason"] = validation_message
                results["files"].append(file_info)
                results["failed_integrations"] += 1
                continue

            try:
                # PROCESAMIENTO REAL DEL ARCHIVO
                file_path = Path(file_info["file_path"])
                content, success = AdvancedFileProcessor.process_file(file_path)

                if not success:
                    print(f"   ⚠️  {content}\n")
                    file_info["status"] = "processing_error"
                    file_info["error"] = content
                    results["files"].append(file_info)
                    results["failed_integrations"] += 1
                    continue

                print(f"   ✓ Contenido extraído: {len(content)} caracteres")

                # Generar ID único
                memory_id = self.generate_memory_id(file_path)

                # Crear entrada de memoria
                memory_entry = self.create_memory_entry(file_info, content, memory_id)
                print(f"   ✓ Chunks creados: {memory_entry['metadata']['total_chunks']}")
                print(f"   ✓ Categoría: {memory_entry['context']['category']}")
                print(f"   ✓ Tags: {', '.join(memory_entry['context']['tags'][:3])}")

                # Guardar memoria
                if self.save_memory(memory_entry):
                    # Mover archivo a procesados
                    dest_path = self.processed_dir / file_path.name
                    shutil.move(str(file_path), str(dest_path))

                    file_info["status"] = "integrated"
                    file_info["memory_id"] = memory_id
                    results["processed_files"] += 1
                    results["successful_integrations"] += 1
                    print(f"   ✅ Integrado exitosamente [ID: {memory_id}]\n")
                else:
                    file_info["status"] = "save_error"
                    results["failed_integrations"] += 1
                    print(f"   ❌ Error guardando memoria\n")

            except Exception as e:
                file_info["status"] = "error"
                file_info["error"] = str(e)
                results["failed_integrations"] += 1
                print(f"   ❌ Error: {e}\n")

            results["files"].append(file_info)

        # Resumen final
        print("=" * 70)
        print("📊 RESUMEN DE INTEGRACIÓN:")
        print(f"   • Total archivos: {results['total_files']}")
        print(f"   • Procesados exitosamente: {results['successful_integrations']}")
        print(f"   • Errores: {results['failed_integrations']}")
        print(f"   • Total memorias en sistema: {self.memory_index['total_memories']}")
        print(f"   • Categorías: {len(self.memory_index['categories'])}")
        print(f"   • Tags únicos: {len(self.memory_index['tags'])}")
        print(f"   • Tipos de archivo: {len(self.memory_index['file_types'])}")

        return results


def print_dependencies_status():
    """Mostrar estado de dependencias"""
    print("\n🔧 ESTADO DE DEPENDENCIAS:")
    print("=" * 70)

    deps = {
        "📄 Documentos": {
            "PyPDF2 (PDF)": PDF_AVAILABLE,
            "python-docx (DOCX)": DOCX_AVAILABLE,
        },
        "📊 Datos": {
            "pandas (Excel/CSV)": PANDAS_AVAILABLE,
            "openpyxl (Excel)": OPENPYXL_AVAILABLE,
        },
        "🖼️  Imágenes": {
            "Pillow (PIL)": PIL_AVAILABLE,
            "pytesseract (OCR)": TESSERACT_AVAILABLE,
        },
        "🎵 Multimedia": {
            "SpeechRecognition (Audio)": SPEECH_AVAILABLE,
            "opencv-python (Video)": CV2_AVAILABLE,
        },
        "🔧 Utilidades": {
            "chardet (Encoding)": CHARDET_AVAILABLE,
        },
    }

    for category, libraries in deps.items():
        print(f"\n{category}")
        for lib, available in libraries.items():
            status = "✅ Instalado" if available else "⚠️  No instalado"
            print(f"   • {lib}: {status}")

    print("\n💡 COMANDOS DE INSTALACIÓN:")
    print("-" * 70)

    if not all([PDF_AVAILABLE, DOCX_AVAILABLE]):
        print("# Documentos:")
        print("pip install PyPDF2 python-docx")

    if not all([PANDAS_AVAILABLE, OPENPYXL_AVAILABLE]):
        print("\n# Datos:")
        print("pip install pandas openpyxl")

    if not all([PIL_AVAILABLE, TESSERACT_AVAILABLE]):
        print("\n# OCR (requiere Tesseract instalado en sistema):")
        print("pip install pillow pytesseract")
        print("# Linux: sudo apt install tesseract-ocr tesseract-ocr-spa")
        print("# Mac: brew install tesseract tesseract-lang")
        print("# Windows: Descargar de https://github.com/UB-Mannheim/tesseract/wiki")

    if not SPEECH_AVAILABLE:
        print("\n# Audio:")
        print("pip install SpeechRecognition pydub")

    if not CV2_AVAILABLE:
        print("\n# Video:")
        print("pip install opencv-python")

    if not CHARDET_AVAILABLE:
        print("\n# Utilidades:")
        print("pip install chardet")

    print("\n# Instalación completa:")
    print("pip install PyPDF2 python-docx pandas openpyxl pillow pytesseract SpeechRecognition opencv-python chardet")
    print("=" * 70)


def main():
    """Función principal"""
    print_dependencies_status()

    integrator = SheilyMemoryIntegrator()
    results = integrator.process_memory_integration()

    if results["total_files"] == 0:
        print("\n💡 INSTRUCCIONES DE USO:")
        print("=" * 70)
        print(f"1. Coloca archivos en: {integrator.input_dir}")
        print("2. Ejecuta: python memory_integrator.py")
        print("3. Los archivos procesados aparecerán en:")
        print(f"   • Originales: {integrator.processed_dir}")
        print(f"   • Memorias: {integrator.memory_dirs['processed']}")
        print(f"   • Índice: {integrator.memory_index_file}")
        print("\n📋 TIPOS DE ARCHIVO SOPORTADOS:")
        print("   • Documentos: PDF, DOCX, TXT, MD, HTML, XML")
        print("   • Datos: CSV, XLSX, XLS, JSON")
        print("   • Código: PY, JS, JAVA, CPP, C, CSS, SQL, etc.")
        print("   • Imágenes: PNG, JPG, JPEG, BMP, TIFF, GIF (con OCR)")
        print("   • Audio: WAV, MP3, OGG, FLAC (con transcripción)")
        print("   • Video: MP4, AVI, MKV, MOV, WMV (metadata)")
        print("   • Archivos: ZIP, TAR, TAR.GZ")
        print("=" * 70)


if __name__ == "__main__":
    main()
